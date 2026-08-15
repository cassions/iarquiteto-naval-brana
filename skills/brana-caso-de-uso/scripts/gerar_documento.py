#!/usr/bin/env python3
"""
Gera o documento final do Caso de Uso a partir de um JSON com as respostas.

Uso:
    python gerar_documento.py respostas.json [-o pasta_de_saida]

O nome do arquivo sai como Caso_de_Uso_<projetista>.docx.
Campos ausentes viram "N/A". Requer python-docx.
"""

import argparse
import json
import re
import sys
import unicodedata
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    sys.exit("python-docx nao encontrado. Rode: pip install python-docx --break-system-packages")

INSTRUCOES = (
    "Despenda algum tempo refletindo sobre os requisitos e uso da embarcacao a ser "
    "desenvolvida. Uma embarcacao e uma solucao de compromisso entre requisitos "
    "conflitantes, este questionario ira nos ajudar a priorizar estas demandas. "
    "Lembre que um barco precisa de uma missao, tarefa ou proposito a ser cumprido, "
    "este sera nosso norte nas escolhas de projeto."
)

# (titulo da secao, [(rotulo, chave), ...])
SECOES = [
    ("Missao", [
        ("Uso principal da embarcacao", "uso_principal"),
        ("Papel mais importante que o barco precisa cumprir", "papel_principal"),
    ]),
    ("Perfil do Usuario", [
        ("Usuario principal", "usuario_principal"),
        ("O usuario principal e tambem o proprietario?", "usuario_e_dono"),
        ("Quem ira pilotar a embarcacao", "quem_pilota"),
        ("Pernoite (havera? para quantas pessoas?)", "pernoite"),
        ("O que mais gosta durante o caminho", "gosta_caminho"),
        ("O que mais gosta enquanto parado", "gosta_parado"),
        ("Outros desejos do usuario ou capitao", "outros_desejos"),
        ("Usuarios em potencial", "usuarios_potenciais"),
        ("Tipo de grupo", "tipo_grupo"),
        ("Numero medio de tripulantes", "tripulantes_medio"),
        ("Numero maximo de tripulantes (dia)", "tripulantes_max_dia"),
        ("Numero maximo de tripulantes (pernoite)", "tripulantes_max_pernoite"),
        ("Local de residencia do usuario principal", "residencia"),
        ("Tipo de casa", "tipo_casa"),
        ("Media anual de dias utilizando a embarcacao", "dias_ano"),
        ("Media anual (dias) de uso continuo", "dias_continuos"),
    ]),
    ("Area de Navegacao", [
        ("Regiao", "regiao"),
        ("Derrota habitual (ponto inicial, ponto final, distancia)", "derrota"),
        ("Pontes e eclusas no trajeto", "pontes_eclusas"),
        ("Clima", "clima"),
        ("Vento / ondas", "vento_ondas"),
        ("Regime de chuvas", "chuvas"),
        ("Variacao de mare", "mare"),
        ("Restricao de calado (calado maximo)", "restricao_calado"),
        ("Outras informacoes relevantes / restricoes", "restricoes_gerais"),
        ("Operacao - tempo", "tempo_operacao"),
        ("Travessias (distancia maxima)", "travessia_max"),
    ]),
    ("Orcamento", [
        ("Orcamento total para construcao/aquisicao", "orcamento_total"),
        ("Capacidade de desembolso mensal", "desembolso_mensal"),
        ("Orcamento anual para manutencao/melhorias", "orcamento_manutencao"),
        ("Mantido por", "mantido_por"),
    ]),
    ("Construtor", [
        ("Construtor / Estaleiro", "construtor"),
        ("Tipo de construtor", "tipo_construtor"),
        ("Experiencia do construtor/estaleiro", "experiencia_construtor"),
        ("Material de construcao", "material"),
        ("Metodo de construcao", "metodo_construcao"),
        ("Faixa de tamanho em que tem experiencia", "faixa_tamanho_experiencia"),
        ("Tipo de embarcacao em que tem experiencia", "tipo_embarcacao_experiencia"),
    ]),
    ("Outros Requisitos", [
        ("Propulsao principal", "propulsao"),
        ("Motorizacao", "motorizacao"),
        ("Armacao (veleiros)", "armacao"),
        ("Requisitos de performance", "performance"),
    ]),
    ("Dimensoes Principais", [
        ("Comprimento (LOA)", "loa"),
        ("Boca (BOA)", "boca"),
        ("Pontal", "pontal"),
        ("Calado", "calado"),
        ("Justificativa das dimensoes", "justificativa_dimensoes"),
    ]),
]

AZUL = RGBColor(0x1F, 0x4E, 0x79)


def slug(texto: str) -> str:
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode()
    texto = re.sub(r"[^A-Za-z0-9]+", "_", texto).strip("_")
    return texto or "Projetista"


def construir(dados: dict, saida: Path) -> Path:
    projetista = dados.get("projetista") or "Projetista"
    doc = Document()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    titulo = doc.add_heading("Caso de Uso", level=0)
    for run in titulo.runs:
        run.font.color.rgb = AZUL

    p = doc.add_paragraph()
    p.add_run("Projetista: ").bold = True
    p.add_run(str(projetista))
    p = doc.add_paragraph()
    p.add_run("Data: ").bold = True
    p.add_run(str(dados.get("data") or date.today().strftime("%d/%m/%Y")))

    h = doc.add_heading("Instrucoes", level=1)
    h.runs[0].font.color.rgb = AZUL
    doc.add_paragraph(INSTRUCOES)

    for secao, campos in SECOES:
        h = doc.add_heading(secao, level=1)
        h.runs[0].font.color.rgb = AZUL
        for rotulo, chave in campos:
            valor = dados.get(chave)
            valor = str(valor).strip() if valor not in (None, "") else "N/A"
            par = doc.add_paragraph()
            r = par.add_run(f"{rotulo}: ")
            r.bold = True
            r.italic = True
            par.add_run(valor)

    saida.mkdir(parents=True, exist_ok=True)
    caminho = saida / f"Caso_de_Uso_{slug(projetista)}.docx"
    doc.save(caminho)
    return caminho


def main() -> None:
    ap = argparse.ArgumentParser(description="Gera o Caso de Uso preenchido em .docx")
    ap.add_argument("json", help="arquivo JSON com as respostas")
    ap.add_argument("-o", "--out", default="/mnt/user-data/outputs", help="pasta de saida")
    args = ap.parse_args()

    dados = json.loads(Path(args.json).read_text(encoding="utf-8"))
    caminho = construir(dados, Path(args.out))
    print(caminho)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
Gera o documento final do Estudo de Similares no modelo Brana.

Uso:
    python gerar_documento.py dados.json [-o pasta_de_saida] \
        [--modelo assets/modelo_tabela_similar_embranco.docx] \
        [--planilha "Estudo de Similar - Nome.xlsx"]

Saida: "Estudo de Similar - <Projetista>.docx", partindo do modelo em branco da
Brana (que ja traz o rodape institucional), preenchendo a tabela comparativa e
acrescentando uma secao por embarcacao com foto, dimensoes e fonte.

Todo valor sai em SI, com duas excecoes que seguem o rotulo do modelo Brana:
potencia em HP e a razao deslocamento/potencia em kg/HP. As conversoes usam os
fatores de scripts/parametros.py - nenhum numero e arredondado na coleta, so na
exibicao (2 casas para metros, 1 para nos, inteiro para kg/L/HP).

Campo que a fonte nao informa fica em branco e aparece na lista de lacunas.
Requer python-docx.
"""

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    sys.exit("python-docx nao encontrado. Rode: pip install python-docx --break-system-packages")

sys.path.insert(0, str(Path(__file__).resolve().parent))
from parametros import FATOR_VALOR, PARAMETROS, POR_KEY, valida_similar  # noqa: E402

VERMELHO = RGBColor(0xC0, 0x00, 0x00)
AZUL = RGBColor(0x1F, 0x38, 0x64)

MODELO_PADRAO = (Path(__file__).resolve().parent.parent / "assets"
                 / "modelo_tabela_similar_embranco.docx")


# ---------------------------------------------------------------- utilidades
def norm(txt):
    txt = unicodedata.normalize("NFKD", str(txt)).encode("ascii", "ignore").decode()
    return re.sub(r"\s+", " ", txt).strip().lower()


def slug(txt):
    txt = unicodedata.normalize("NFKD", str(txt)).encode("ascii", "ignore").decode()
    txt = re.sub(r"[^\w\s-]", "", txt).strip()
    return re.sub(r"\s+", " ", txt) or "Projetista"


def num(v, dec=2, milhar=False):
    """Formata numero no padrao pt-BR. Devolve '' para ausente."""
    if v is None or v == "":
        return ""
    if isinstance(v, str):
        return v
    if milhar:
        s = f"{v:,.{dec}f}".replace(",", " ").replace(".", ",")
        return s.replace(" ", ".")
    return f"{v:.{dec}f}".replace(".", ",")


def nome_completo(sim):
    est = (sim.get("estaleiro") or "").strip()
    nome = (sim.get("nome") or "").strip()
    ano = (str(sim.get("ano") or "")).strip()
    txt = f"{est} {nome}".strip()
    return f"{txt} ({ano})" if ano else txt


def nome_curto(sim):
    est = (sim.get("estaleiro") or "").strip()
    nome = (sim.get("nome") or "").strip()
    return f"{est} {nome}".strip() or "(sem nome)"


def em_si(key, valor, sistema):
    """Converte o valor publicado para SI, se a fonte era imperial."""
    p = POR_KEY.get(key)
    if valor is None or p is None or p["tipo"] == "texto":
        return valor
    if sistema == "imperial" and p["fator"]:
        return valor * FATOR_VALOR[p["fator"]]
    return valor


def potencia_hp(sim):
    """Potencia em HP - o modelo Brana rotula essa linha em HP."""
    v = (sim.get("dados") or {}).get("potencia")
    if v is None:
        return None
    return v if sim.get("sistema") == "imperial" else v / FATOR_VALOR["HP_KW"]


# ------------------------------------------------- valores para a tabela docx
def valor_docx(key, sim):
    """Devolve a string ja formatada para a celula da tabela do modelo."""
    d = sim.get("dados") or {}
    sistema = sim.get("sistema")
    pub = d.get(key)

    if key == "potencia":
        return num(potencia_hp(sim), 0, milhar=True)
    if key == "motorizacao":
        return (pub or "").strip()
    if key in ("loa_boa", "dlr", "disp_potencia"):
        if pub is not None:
            return num(pub, 0 if key == "dlr" else 2)
        loa = em_si("loa", d.get("loa"), sistema)
        boa = em_si("boa", d.get("boa"), sistema)
        lwl = em_si("lwl", d.get("lwl"), sistema)
        disp = em_si("deslocamento", d.get("deslocamento"), sistema)
        hp = potencia_hp(sim)
        if key == "loa_boa":
            return num(loa / boa, 2) if loa and boa else ""
        if key == "dlr" and disp and lwl:
            lwl_ft = lwl / FATOR_VALOR["FT_M"]
            return num((disp / FATOR_VALOR["LT_KG"]) / ((0.01 * lwl_ft) ** 3), 0)
        if key == "disp_potencia" and disp and hp:
            return num(disp / hp, 1)
        return ""

    v = em_si(key, pub, sistema)
    if v is None:
        return ""
    p = POR_KEY[key]
    if p["fator"] in ("LB_KG", "GAL_L"):
        return num(v, 0, milhar=True)
    if p["fator"] == "MPH_KN":
        return num(v, 1)
    if p["key"] in ("deadrise", "autonomia"):
        return num(v, 0, milhar=True)
    return num(v, 2)


# rotulo do modelo (normalizado) -> chave dos dados
MAPA_TABELA = {
    "comprimento total (loa)": "loa",
    "comprimento na linha d'agua (lwl)": "lwl",
    "boca maxima (boa)": "boa",
    "calado": "calado",
    "deslocamento": "deslocamento",
    "angulo de v do fundo (deadrise)": "deadrise",
    "potencia instalada": "potencia",
    "motorizacao (marca / modelo / qtd.)": "motorizacao",
    "area velica (se veleiro)": "area_velica",
    "combustivel": "combustivel",
    "agua doce": "agua",
    "velocidade maxima": "vel_max",
    "velocidade de cruzeiro": "vel_cruzeiro",
    "autonomia": "autonomia",
    "loa / boa": "loa_boa",
    "razao desloc.-comprimento (dlr)": "dlr",
    "deslocamento / potencia": "disp_potencia",
}

# parametros que nao estao na tabela do modelo mas podem ter sido coletados
EXTRAS = ["pontal", "pe_direito", "dejetos", "ff", "fm", "para_brisa",
          "deck_cl", "plataforma_wl", "balaustre", "motorizacao_det"]


# ------------------------------------------------------------------ paragrafos
def par(doc, texto="", tam=10, negrito=False, italico=False, cor=None,
        espaco_antes=0, espaco_depois=4, alinhamento=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(espaco_antes)
    p.paragraph_format.space_after = Pt(espaco_depois)
    if alinhamento is not None:
        p.alignment = alinhamento
    if texto:
        r = p.add_run(texto)
        r.font.size = Pt(tam)
        r.bold = negrito
        r.italic = italico
        if cor is not None:
            r.font.color.rgb = cor
    return p


def mover_antes(paragrafo, tbl):
    """Move um paragrafo (criado no fim do documento) para antes da tabela."""
    tbl._tbl.addprevious(paragrafo._p)


def escreve_celula(cel, texto, negrito=False, tam=8.5, centro=True):
    cel.text = ""
    p = cel.paragraphs[0]
    if centro:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(texto))
    r.font.size = Pt(tam)
    r.bold = negrito


# --------------------------------------------------------------------- tabela
def preenche_tabela(tbl, similares):
    n = len(similares)
    n_col_barcos = len(tbl.columns) - 2
    while n_col_barcos < n:                       # mais similares que o modelo
        tbl.add_column(Cm(2.4))
        n_col_barcos += 1

    # cabecalho: numero -> nome curto
    for i in range(n_col_barcos):
        cel = tbl.rows[0].cells[2 + i]
        escreve_celula(cel, nome_curto(similares[i]) if i < n else "", negrito=True)

    for row in tbl.rows[1:]:
        rot = norm(row.cells[0].text)
        if not rot:
            continue
        if rot.startswith("embarcacao"):
            for i in range(n):
                escreve_celula(row.cells[2 + i], nome_completo(similares[i]),
                               negrito=True)
            continue
        if rot.startswith("razoes"):
            continue
        key = MAPA_TABELA.get(rot)
        if key is None:
            continue
        for i in range(n):
            texto = valor_docx(key, similares[i])
            escreve_celula(row.cells[2 + i], texto,
                           centro=(key != "motorizacao"))


# --------------------------------------------------------------------- secoes
def secao_similar(doc, sim, tbl, indice, base_dir):
    d = sim.get("dados") or {}
    sistema = sim.get("sistema")

    p = par(doc, f"Similar {indice} — {nome_completo(sim)}", tam=12,
            negrito=True, cor=VERMELHO, espaco_antes=10, espaco_depois=2)
    mover_antes(p, tbl)

    if sim.get("porque_similar"):
        p = par(doc, f"Por que serve de referência: {sim['porque_similar']}", tam=10)
        mover_antes(p, tbl)

    # foto
    foto = sim.get("foto")
    if foto:
        cam = Path(foto)
        if not cam.is_absolute():
            cam = base_dir / cam
        if cam.exists():
            try:
                doc.add_picture(str(cam), width=Cm(11))
                pfoto = doc.paragraphs[-1]
                pfoto.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mover_antes(pfoto, tbl)
                cred = sim.get("foto_credito") or ""
                p = par(doc, f"Foto: {cred}" if cred else "Foto sem crédito informado.",
                        tam=8, italico=True, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
                mover_antes(p, tbl)
            except Exception as e:                       # imagem invalida
                p = par(doc, f"[imagem nao inserida: {e}]", tam=8, italico=True)
                mover_antes(p, tbl)
        else:
            p = par(doc, f"[foto nao encontrada: {foto}]", tam=8, italico=True)
            mover_antes(p, tbl)

    # dimensoes principais
    principais = [("LOA", "loa", "m"), ("LWL", "lwl", "m"), ("BOA (boca)", "boa", "m"),
                  ("Calado", "calado", "m"), ("Pontal", "pontal", "m"),
                  ("Deslocamento", "deslocamento", "kg")]
    partes = []
    for rot, key, un in principais:
        txt = valor_docx(key, sim) if key in MAPA_TABELA.values() else \
            num(em_si(key, d.get(key), sistema), 2)
        if txt:
            partes.append(f"{rot}: {txt} {un}")
    hp = potencia_hp(sim)
    if hp:
        partes.append(f"Potência: {num(hp, 0, milhar=True)} HP")
    if d.get("motorizacao"):
        partes.append(f"Motorização: {d['motorizacao']}")
    if partes:
        p = par(doc, "  •  ".join(partes), tam=10)
        mover_antes(p, tbl)

    # extras coletados que nao entram na tabela do modelo
    ex = []
    for key in EXTRAS:
        v = d.get(key)
        if v in (None, ""):
            continue
        p_ = POR_KEY[key]
        if p_["tipo"] == "texto":
            ex.append(f"{p_['pt']}: {v}")
        else:
            ex.append(f"{p_['pt']}: {num(em_si(key, v, sistema), 2)} {p_['unid_si']}")
    if ex:
        p = par(doc, "Outros dados publicados — " + "; ".join(ex), tam=9)
        mover_antes(p, tbl)

    # fontes
    for f in sim.get("fontes") or []:
        data = f.get("data_publicacao") or "sem data identificável"
        acesso = f" | acesso em {f['data_acesso']}" if f.get("data_acesso") else ""
        tipo = f" [{f['tipo']}]" if f.get("tipo") else ""
        p = par(doc, f"Fonte{tipo}: {f.get('nome', '')} — {data}{acesso} — "
                     f"{f.get('url', '')}", tam=8.5, italico=True)
        mover_antes(p, tbl)

    if sim.get("lacunas"):
        rotulos = [POR_KEY[k]["pt"] if k in POR_KEY else k for k in sim["lacunas"]]
        p = par(doc, "Lacunas (a fonte não informa): " + ", ".join(rotulos), tam=8.5,
                italico=True)
        mover_antes(p, tbl)

    if sim.get("divergencias"):
        p = par(doc, f"Divergência entre fontes: {sim['divergencias']}", tam=8.5,
                italico=True, cor=VERMELHO)
        mover_antes(p, tbl)


# ----------------------------------------------------------------------- main
def gerar(dados, saida, modelo, planilha=None, base_dir=Path(".")):
    similares = dados.get("similares") or []
    if len(similares) < 3:
        print(f"AVISO: apenas {len(similares)} similar(es); o estudo pede no minimo 3.",
              file=sys.stderr)
    erros = []
    for i, sim in enumerate(similares):
        erros += valida_similar(sim, i)
    if erros:
        print("Problemas nos dados:\n  - " + "\n  - ".join(erros), file=sys.stderr)
        sys.exit(1)

    doc = Document(str(modelo))
    projetista = dados.get("projetista") or ""

    # ---- titulo e subtitulos do modelo
    if doc.paragraphs:
        p0 = doc.paragraphs[0]
        for r in p0.runs[1:]:
            r.text = ""
        if p0.runs:
            p0.runs[0].text = f"ESTUDO DE SIMILAR — {projetista.upper()}"
        else:
            p0.add_run(f"ESTUDO DE SIMILAR — {projetista.upper()}")
    # o modelo em branco diz "preencher a mao"; aqui a tabela ja vem preenchida
    if len(doc.paragraphs) > 1:
        p1 = doc.paragraphs[1]
        if "preencher" in norm(p1.text):
            for r in p1.runs[1:]:
                r.text = ""
            if p1.runs:
                p1.runs[0].text = ("Estudo comparativo de embarcações similares  |  "
                                   "unidades no sistema métrico (SI)")

    tbl = doc.tables[0]

    def antes(texto, **kw):
        mover_antes(par(doc, texto, **kw), tbl)

    antes(f"Projetista: {projetista}", tam=10, negrito=True, espaco_antes=8,
          espaco_depois=1)
    if dados.get("data"):
        antes(f"Data: {dados['data']}", tam=10, espaco_depois=1)
    if dados.get("embarcacao_pretendida"):
        antes(f"Embarcação pretendida: {dados['embarcacao_pretendida']}", tam=10)

    if dados.get("caso_de_uso_resumo"):
        antes("1. Caso de uso", tam=12, negrito=True, cor=AZUL, espaco_antes=10,
              espaco_depois=2)
        antes(dados["caso_de_uso_resumo"], tam=10)

    if dados.get("criterios_selecao"):
        antes("2. Critérios de seleção dos similares", tam=12, negrito=True,
              cor=AZUL, espaco_antes=10, espaco_depois=2)
        antes(dados["criterios_selecao"], tam=10)

    antes("3. Embarcações similares", tam=12, negrito=True, cor=AZUL,
          espaco_antes=10, espaco_depois=2)
    for i, sim in enumerate(similares, start=1):
        secao_similar(doc, sim, tbl, i, base_dir)

    antes("4. Tabela comparativa", tam=12, negrito=True, cor=AZUL,
          espaco_antes=12, espaco_depois=2)
    antes("Valores em SI, exceto potência (HP) e deslocamento/potência (kg/HP), "
          "que seguem o rótulo deste modelo. Célula em branco significa que a "
          "fonte não informa o dado — nunca que o valor é zero.", tam=8.5,
          italico=True)

    preenche_tabela(tbl, similares)

    # ---- conteudo depois da tabela
    if dados.get("leitura_parametrica"):
        par(doc, "5. Leitura paramétrica", tam=12, negrito=True, cor=AZUL,
            espaco_antes=12, espaco_depois=2)
        par(doc, dados["leitura_parametrica"], tam=10)

    if dados.get("recomendacao"):
        par(doc, "6. Recomendação preliminar para a embarcação nova", tam=12,
            negrito=True, cor=AZUL, espaco_antes=10, espaco_depois=2)
        par(doc, dados["recomendacao"], tam=10)

    par(doc, "7. Lacunas e observações", tam=12, negrito=True, cor=AZUL,
        espaco_antes=10, espaco_depois=2)
    achou = False
    for sim in similares:
        itens = []
        if sim.get("lacunas"):
            rot = [POR_KEY[k]["pt"] if k in POR_KEY else k for k in sim["lacunas"]]
            itens.append("lacunas: " + ", ".join(rot))
        if sim.get("divergencias"):
            itens.append("divergência: " + sim["divergencias"])
        if itens:
            par(doc, f"• {nome_completo(sim)} — " + " | ".join(itens), tam=9.5)
            achou = True
    if not achou:
        par(doc, "Nenhuma lacuna ou divergência registrada.", tam=9.5, italico=True)
    if dados.get("observacoes"):
        par(doc, dados["observacoes"], tam=10, espaco_antes=4)

    par(doc, "8. Fontes consultadas", tam=12, negrito=True, cor=AZUL,
        espaco_antes=10, espaco_depois=2)
    for sim in similares:
        for f in sim.get("fontes") or []:
            data = f.get("data_publicacao") or "sem data identificável"
            acesso = f" | acesso em {f['data_acesso']}" if f.get("data_acesso") else ""
            par(doc, f"• {nome_curto(sim)} — {f.get('nome', '')} — {data}{acesso} — "
                     f"{f.get('url', '')}", tam=9)

    if planilha:
        par(doc, f"Planilha comparativa completa (todos os parâmetros, aba com o "
                 f"dado original da fonte e aba métrica calculada por fórmula): "
                 f"{Path(planilha).name}", tam=9, italico=True, espaco_antes=8)

    saida = Path(saida)
    saida.mkdir(parents=True, exist_ok=True)
    arq = saida / f"Estudo de Similar - {slug(projetista)}.docx"
    doc.save(arq)
    return arq


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("json", help="arquivo dados.json")
    ap.add_argument("-o", "--out", default=".", help="pasta de saida")
    ap.add_argument("--modelo", default=str(MODELO_PADRAO),
                    help="modelo .docx em branco da Brana")
    ap.add_argument("--planilha", default=None,
                    help="nome do .xlsx gerado, para citar no documento")
    a = ap.parse_args()
    caminho = Path(a.json)
    dados = json.loads(caminho.read_text(encoding="utf-8"))
    print(gerar(dados, a.out, a.modelo, a.planilha, caminho.resolve().parent))


if __name__ == "__main__":
    main()
